from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "word"
IMG_OUT = ROOT / "generated_images"

TITLE_FONT = "微软雅黑"
BODY_FONT = "宋体"


def _font_path(name: str) -> str:
    candidates = {
        "宋体": [
            r"C:\Windows\Fonts\simsun.ttc",
            r"C:\Windows\Fonts\simsun.ttf",
        ],
        "微软雅黑": [
            r"C:\Windows\Fonts\msyh.ttc",
            r"C:\Windows\Fonts\msyh.ttf",
        ],
    }
    for path in candidates.get(name, []):
        if Path(path).exists():
            return path
    return r"C:\Windows\Fonts\simsun.ttc"


def _pil_font(size: int, *, bold: bool = False) -> ImageFont.FreeTypeFont:
    font_path = _font_path("微软雅黑" if bold else "宋体")
    return ImageFont.truetype(font_path, size=size)


def _set_run_font(run, *, name: str = BODY_FONT, size: int = 12, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def _set_paragraph_spacing(paragraph, *, before: int = 0, after: int = 6, line: float = 1.5) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def _configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(12)


def _add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p, after=12, line=1.2)
    run = p.add_run(text.strip())
    _set_run_font(run, name=TITLE_FONT, size=18, bold=True)


def _add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=8 if level <= 2 else 4, after=6, line=1.3)
    if level <= 2:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        size = 15
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        size = 13
    run = p.add_run(text.strip())
    _set_run_font(run, name=TITLE_FONT, size=size, bold=True)


def _add_body_paragraph(doc: Document, text: str, *, indent: bool = True) -> None:
    p = doc.add_paragraph()
    _set_paragraph_spacing(p)
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text.strip())
    _set_run_font(run)


def _add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style=None)
    _set_paragraph_spacing(p, after=3)
    p.paragraph_format.left_indent = Pt(24)
    p.paragraph_format.first_line_indent = Pt(-12)
    run = p.add_run("· " + text.strip())
    _set_run_font(run)


def _add_markdown(doc: Document, path: Path, *, skip_top_title: bool = False) -> None:
    lines = path.read_text(encoding="utf-8").splitlines()
    in_code = False
    table_buffer: list[str] = []

    def flush_table() -> None:
        nonlocal table_buffer
        if not table_buffer:
            return
        # Keep simple markdown tables readable without complex styling.
        for row in table_buffer:
            if re.match(r"^\|\s*-", row):
                continue
            cells = [cell.strip() for cell in row.strip("|").split("|")]
            _add_body_paragraph(doc, " | ".join(cells), indent=False)
        table_buffer = []

    for raw in lines:
        line = raw.rstrip()
        if line.strip().startswith("```"):
            flush_table()
            in_code = not in_code
            continue
        if in_code:
            if line.strip():
                _add_body_paragraph(doc, line, indent=False)
            continue
        if not line.strip():
            flush_table()
            continue
        if line.startswith("|"):
            table_buffer.append(line)
            continue
        flush_table()
        if line.startswith("# "):
            if not skip_top_title:
                _add_title(doc, line[2:])
            continue
        if line.startswith("## "):
            _add_heading(doc, line[3:], 2)
            continue
        if line.startswith("### "):
            _add_heading(doc, line[4:], 3)
            continue
        if line.startswith("- "):
            _add_bullet(doc, line[2:])
            continue
        _add_body_paragraph(doc, line, indent=not re.match(r"^\d+\. ", line))
    flush_table()


def _arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int]) -> None:
    draw.line([start, end], fill="black", width=3)
    x1, y1 = start
    x2, y2 = end
    if abs(x2 - x1) >= abs(y2 - y1):
        direction = 1 if x2 >= x1 else -1
        points = [(x2, y2), (x2 - 18 * direction, y2 - 9), (x2 - 18 * direction, y2 + 9)]
    else:
        direction = 1 if y2 >= y1 else -1
        points = [(x2, y2), (x2 - 9, y2 - 18 * direction), (x2 + 9, y2 - 18 * direction)]
    draw.polygon(points, fill="black")


def _center_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], lines: list[str], font) -> None:
    x1, y1, x2, y2 = box
    line_heights = []
    widths = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        widths.append(bbox[2] - bbox[0])
        line_heights.append(bbox[3] - bbox[1])
    total_h = sum(line_heights) + 12 * (len(lines) - 1)
    y = y1 + ((y2 - y1 - total_h) / 2)
    for line, w, h in zip(lines, widths, line_heights):
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, fill="black", font=font)
        y += h + 12


def _box(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], lines: list[str], *, font=None) -> None:
    draw.rectangle(xy, outline="black", width=3)
    _center_text(draw, xy, lines, font or _pil_font(34))


def _make_canvas(title: str) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", (1600, 1100), "white")
    draw = ImageDraw.Draw(img)
    title_font = _pil_font(42, bold=True)
    _center_text(draw, (0, 30, 1600, 90), [title], title_font)
    return img, draw


def _save(img: Image.Image, name: str) -> Path:
    IMG_OUT.mkdir(parents=True, exist_ok=True)
    path = IMG_OUT / name
    img.save(path, "PNG")
    return path


FIGURE_CAPTIONS = [
    "图1",
    "图2",
    "图3",
    "图4",
    "图5",
    "图6",
    "图7",
    "图8",
    "图9",
]


def draw_figures() -> list[Path]:
    images: list[Path] = []
    font = _pil_font(34)
    small = _pil_font(27)
    tiny = _pil_font(23)

    img, d = _make_canvas("图1")
    _box(d, (560, 110, 1040, 205), ["100 内容对象输入"], font=font)
    _arrow(d, (800, 205), (800, 275))
    _box(d, (560, 275, 1040, 370), ["110 字段归一化"], font=font)
    _arrow(d, (800, 370), (800, 440))
    _box(d, (560, 440, 1040, 535), ["120 压缩语义投影"], font=font)
    _arrow(d, (800, 535), (800, 605))
    _box(d, (180, 605, 680, 710), ["130 卡组结构分析", "135 曲线分析"], font=small)
    _box(d, (920, 605, 1420, 710), ["140 敌人压力分析"], font=font)
    _arrow(d, (430, 710), (690, 825))
    _arrow(d, (1170, 710), (910, 825))
    _box(d, (560, 825, 1040, 920), ["150 候选差分评估"], font=font)
    _arrow(d, (800, 920), (800, 990))
    _box(d, (560, 990, 1040, 1070), ["180 审计输出"], font=font)
    images.append(_save(img, "fig1_overall_flow.png"))

    img, d = _make_canvas("图2")
    _box(d, (80, 150, 390, 280), ["卡牌对象", "费用/效果/触发/目标"], font=small)
    _arrow(d, (390, 215), (525, 215))
    _box(d, (525, 130, 875, 300), ["110 字段归一化", "效果单元拆解"], font=small)
    _arrow(d, (875, 215), (1010, 215))
    _box(d, (1010, 150, 1460, 280), ["效果单元", "类型/目标/数值", "条件/窗口/区域"], font=small)
    _arrow(d, (1235, 280), (1235, 390))
    _box(d, (935, 390, 1535, 510), ["初始特征元组", "效果类型+数值幅度+触发条件"], font=small)
    _arrow(d, (935, 450), (690, 590))
    _box(d, (310, 555, 860, 685), ["120 压缩语义投影", "能力轴/状态/角色/关系字典映射"], font=small)
    _arrow(d, (585, 685), (585, 795))
    d.rectangle((90, 795, 1080, 1040), outline="black", width=3)
    _center_text(d, (90, 810, 1080, 865), ["压缩语义表示"], font)
    for xy, text in [
        ((145, 890, 390, 950), "成本向量"),
        ((465, 890, 710, 950), "收益向量"),
        ((785, 890, 1030, 950), "状态亲和"),
        ((145, 970, 390, 1030), "角色标签"),
        ((465, 970, 710, 1030), "关系边"),
        ((785, 970, 1030, 1030), "不确定性"),
    ]:
        _box(d, xy, [text], font=small)
    _arrow(d, (1080, 910), (1240, 910))
    _box(d, (1240, 840, 1510, 980), ["归一化/饱和", "置信度标注"], font=small)
    images.append(_save(img, "fig2_effect_unit_projection.png"))

    img, d = _make_canvas("图3")
    d.rectangle((120, 135, 1480, 1015), outline="black", width=3)
    _center_text(d, (120, 150, 1480, 220), ["压缩语义表示的数据结构"], font)
    boxes = [
        ((190, 280, 555, 390), ["成本向量", "能量/手牌/节奏", "条件/风险"]),
        ((625, 280, 990, 390), ["收益向量", "输出/生存/资源", "引擎/选择/控制"]),
        ((1060, 280, 1425, 390), ["状态亲和向量", "启动/过渡/在线", "转化/恢复"]),
        ((190, 520, 555, 630), ["角色标签集合", "启动件/桥接件", "收益件/稳定件"]),
        ((625, 520, 990, 630), ["机制关系边", "生产/消耗/放大", "转化/依赖/稳定"]),
        ((1060, 520, 1425, 630), ["不确定性标签", "未映射效果", "高方差/投影缺口"]),
    ]
    for xy, lines in boxes:
        _box(d, xy, lines, font=tiny)
    _box(d, (445, 800, 1155, 930), ["统一比较空间", "卡牌对象 / 卡组对象 / 敌人对象 / 候选对象"], font=small)
    for start in [(370, 390), (800, 390), (1235, 390), (370, 630), (800, 630), (1235, 630)]:
        _arrow(d, start, (800, 800))
    images.append(_save(img, "fig3_semantic_structure.png"))

    img, d = _make_canvas("图4")
    _box(d, (95, 150, 425, 260), ["多张卡牌", "压缩语义表示"], font=small)
    _arrow(d, (425, 205), (555, 205))
    _box(d, (555, 130, 1045, 280), ["130 卡组结构分析", "关系边聚合 / 角色覆盖 / 状态覆盖"], font=small)
    _arrow(d, (1045, 205), (1175, 205))
    _box(d, (1175, 150, 1510, 260), ["卡组结构表示"], font=small)
    for xy, lines in [
        ((120, 430, 470, 540), ["机制包权重", "关系边连通性"]),
        ((625, 430, 975, 540), ["能力轴权重", "收益向量聚合"]),
        ((1130, 430, 1480, 540), ["状态/角色覆盖", "阶段与构筑职责"]),
        ((375, 705, 725, 815), ["资源压力", "费用/手牌/资源"]),
        ((875, 705, 1225, 815), ["死牌压力", "条件过窄/时机错位"]),
    ]:
        _box(d, xy, lines, font=small)
    _arrow(d, (795, 280), (295, 430))
    _arrow(d, (795, 280), (800, 430))
    _arrow(d, (795, 280), (1305, 430))
    _arrow(d, (795, 280), (550, 705))
    _arrow(d, (795, 280), (1050, 705))
    _box(d, (530, 910, 1070, 1015), ["结构缺口表示", "缺角色 / 缺状态 / 资源压力 / 机制离散"], font=small)
    _arrow(d, (550, 815), (700, 910))
    _arrow(d, (1050, 815), (900, 910))
    images.append(_save(img, "fig4_deck_structure.png"))

    img, d = _make_canvas("图5")
    _box(d, (130, 135, 470, 235), ["卡组结构表示"], font=small)
    _arrow(d, (470, 185), (610, 185))
    _box(d, (610, 120, 990, 250), ["135 曲线分析", "按费用带/阶段/压力窗口分桶"], font=small)
    _arrow(d, (990, 185), (1130, 185))
    _box(d, (1130, 135, 1470, 235), ["曲线表示"], font=small)
    y_positions = [345, 455, 565, 675, 785, 895]
    labels = ["费用曲线", "启动曲线", "资源曲线", "手牌流转曲线", "收益依赖曲线", "压力答案曲线"]
    for y, label in zip(y_positions, labels):
        _box(d, (180, y, 475, y + 70), [label], font=small)
        d.line([(550, y + 55), (670, y + 20), (790, y + 35), (910, y + 5), (1030, y + 45), (1150, y + 25), (1270, y + 55), (1390, y + 15)], fill="black", width=3)
        d.rectangle((535, y, 1410, y + 70), outline="black", width=2)
    _box(d, (560, 1000, 1040, 1070), ["曲线缺口：过高/不足/断裂/时机错位"], font=small)
    images.append(_save(img, "fig5_deck_curves.png"))

    img, d = _make_canvas("图6")
    _box(d, (110, 145, 610, 265), ["卡组答案画像", "输出/生存/资源/选择/控制"], font=small)
    _box(d, (990, 145, 1490, 265), ["敌人压力语义", "压力向量/干扰/时机窗口"], font=small)
    _arrow(d, (360, 265), (660, 450))
    _arrow(d, (1240, 265), (940, 450))
    _box(d, (560, 450, 1040, 570), ["同一能力轴字典", "供给与需求对齐"], font=small)
    _arrow(d, (800, 570), (800, 710))
    _box(d, (510, 710, 1090, 840), ["压力匹配结果", "缺失答案 / 薄弱答案 / 时机风险"], font=small)
    _arrow(d, (800, 840), (800, 960))
    _box(d, (510, 960, 1090, 1060), ["进入候选贡献计算"], font=small)
    images.append(_save(img, "fig6_pressure_matching.png"))

    img, d = _make_canvas("图7")
    _box(d, (110, 145, 470, 255), ["候选内容对象"], font=small)
    _box(d, (1130, 145, 1490, 255), ["目标卡组/上下文"], font=small)
    _arrow(d, (470, 200), (650, 330))
    _arrow(d, (1130, 200), (950, 330))
    _box(d, (560, 330, 1040, 450), ["临时加入候选", "重新计算结构与曲线"], font=small)
    _arrow(d, (800, 450), (800, 570))
    _box(d, (215, 570, 690, 700), ["加入前", "结构表示 / 曲线表示 / 缺口表示"], font=small)
    _box(d, (910, 570, 1385, 700), ["加入后", "结构表示 / 曲线表示 / 缺口表示"], font=small)
    _arrow(d, (690, 635), (910, 635))
    _box(d, (550, 780, 1050, 900), ["差分计算", "贡献变化量 + 风险变化量"], font=small)
    _arrow(d, (800, 900), (800, 1010))
    _box(d, (455, 1010, 1145, 1080), ["候选贡献表示 / 风险表示 / 候选排序"], font=small)
    images.append(_save(img, "fig7_candidate_delta.png"))

    img, d = _make_canvas("图8")
    _box(d, (220, 130, 1380, 260), ["人工审查样本", "正样本 / 负样本 / 近邻样本 / 误报样本"], font=small)
    _arrow(d, (800, 260), (800, 380))
    _box(d, (520, 380, 1080, 500), ["160 样本构造模块"], font=font)
    _arrow(d, (800, 500), (800, 620))
    _box(d, (220, 620, 1380, 760), ["成对排序样本", "赢家对象 / 输家对象 / 标签类型 / 差异特征"], font=small)
    _arrow(d, (800, 760), (800, 890))
    _box(d, (520, 890, 1080, 1010), ["170 学习排序模块"], font=font)
    images.append(_save(img, "fig8_pairwise_samples.png"))

    img, d = _make_canvas("图9")
    _box(d, (130, 150, 560, 270), ["候选生成"], font=font)
    _box(d, (1040, 150, 1470, 270), ["170 学习排序"], font=font)
    _arrow(d, (560, 210), (1040, 210))
    _arrow(d, (345, 270), (345, 425))
    _box(d, (130, 425, 560, 545), ["基础规则排序"], font=font)
    _arrow(d, (560, 485), (695, 650))
    _arrow(d, (1255, 270), (905, 650))
    _box(d, (560, 650, 1040, 770), ["190 人工审核门控"], font=font)
    _center_text(d, (560, 790, 1040, 850), ["预设验证条件 / 人工审核条件"], small)
    _arrow(d, (800, 770), (800, 900))
    _box(d, (560, 900, 1040, 1020), ["180 审计输出"], font=font)
    images.append(_save(img, "fig9_audit_gate.png"))
    return images


def add_single_drawing(doc: Document, image_path: Path, caption: str, *, page_break: bool = False) -> None:
    if page_break:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(6.5))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    _set_run_font(r, size=12)


def add_drawings(doc: Document, image_paths: list[Path]) -> None:
    _add_title(doc, "说明书附图")
    for index, image_path in enumerate(image_paths):
        add_single_drawing(doc, image_path, FIGURE_CAPTIONS[index], page_break=bool(index))


def save_doc(doc: Document, path: Path) -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def build_text_doc(markdown_path: Path, output_name: str) -> Path:
    doc = Document()
    _configure_document(doc)
    _add_markdown(doc, markdown_path)
    path = OUT / output_name
    save_doc(doc, path)
    return path


def build_preview_spec_doc(markdown_path: Path, image_paths: list[Path]) -> Path:
    doc = Document()
    _configure_document(doc)
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("# "):
            _add_title(doc, line[2:])
        elif line.startswith("## "):
            _add_heading(doc, line[3:], 2)
        elif line.startswith("### "):
            _add_heading(doc, line[4:], 3)
        else:
            _add_body_paragraph(doc, line, indent=not re.match(r"^\d+\. ", line))
            match = re.match(r"^\[\d{4}\]\s*图(\d+)为", line)
            if match:
                figure_index = int(match.group(1)) - 1
                if 0 <= figure_index < len(image_paths):
                    add_single_drawing(doc, image_paths[figure_index], FIGURE_CAPTIONS[figure_index])
    path = OUT / "02_说明书_带附图预览.docx"
    save_doc(doc, path)
    return path


def build_drawings_doc(image_paths: list[Path]) -> Path:
    doc = Document()
    _configure_document(doc)
    add_drawings(doc, image_paths)
    path = OUT / "04_说明书附图.docx"
    save_doc(doc, path)
    return path


def build_reading_doc(image_paths: list[Path]) -> Path:
    doc = Document()
    _configure_document(doc)
    _add_markdown(doc, ROOT / "01_权利要求书.md")
    doc.add_page_break()
    preview_path = ROOT / "02_说明书.md"
    lines = preview_path.read_text(encoding="utf-8").splitlines()
    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("# "):
            _add_title(doc, line[2:])
        elif line.startswith("## "):
            _add_heading(doc, line[3:], 2)
        elif line.startswith("### "):
            _add_heading(doc, line[4:], 3)
        else:
            _add_body_paragraph(doc, line, indent=not re.match(r"^\d+\. ", line))
            match = re.match(r"^\[\d{4}\]\s*图(\d+)为", line)
            if match:
                figure_index = int(match.group(1)) - 1
                if 0 <= figure_index < len(image_paths):
                    add_single_drawing(doc, image_paths[figure_index], FIGURE_CAPTIONS[figure_index])
    doc.add_page_break()
    _add_markdown(doc, ROOT / "03_说明书摘要.md")
    path = OUT / "CardAnalysis_底座发明专利_图文阅读版.docx"
    save_doc(doc, path)
    return path


def build_combined_doc(image_paths: list[Path]) -> Path:
    doc = Document()
    _configure_document(doc)
    files = [
        ROOT / "01_权利要求书.md",
        ROOT / "02_说明书.md",
        ROOT / "03_说明书摘要.md",
    ]
    for index, path in enumerate(files):
        if index:
            doc.add_page_break()
        _add_markdown(doc, path)
    doc.add_page_break()
    add_drawings(doc, image_paths)
    path = OUT / "CardAnalysis_底座发明专利_正式版合并稿.docx"
    save_doc(doc, path)
    return path


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    IMG_OUT.mkdir(parents=True, exist_ok=True)
    for old_image in IMG_OUT.glob("fig*.png"):
        old_image.unlink()
    image_paths = draw_figures()
    outputs = [
        build_text_doc(ROOT / "01_权利要求书.md", "01_权利要求书.docx"),
        build_text_doc(ROOT / "02_说明书.md", "02_说明书.docx"),
        build_preview_spec_doc(ROOT / "02_说明书.md", image_paths),
        build_text_doc(ROOT / "03_说明书摘要.md", "03_说明书摘要.docx"),
        build_drawings_doc(image_paths),
        build_combined_doc(image_paths),
        build_reading_doc(image_paths),
    ]
    for output in outputs:
        print(output)


if __name__ == "__main__":
    main()
